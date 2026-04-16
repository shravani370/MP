"""
Async tasks for resume parsing and processing
"""
from celery import shared_task
from celery.utils.log import get_task_logger
import io

logger = get_task_logger(__name__)

# ═══════════════════════════════════════════════════════════════════════════
# RESUME PARSING TASK
# ═══════════════════════════════════════════════════════════════════════════

@shared_task(bind=True, max_retries=2)
def parse_resume_async(self, file_path: str, user_email: str):
    """
    Async task: Parse resume file (PDF or DOCX)
    Extracts text for AI analysis
    
    Args:
        file_path: Path to uploaded resume file
        user_email: User's email
    
    Returns:
        dict: {"status": "success", "text": "...", "pages": 3, "task_id": "..."}
    """
    try:
        import os
        from PyPDF2 import PdfReader
        from docx import Document
        
        logger.info(f"Parsing resume for {user_email}: {file_path}")
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        resume_text = ""
        page_count = 0
        
        # Parse based on file extension
        if file_path.endswith('.pdf'):
            try:
                with open(file_path, 'rb') as f:
                    pdf = PdfReader(f)
                    page_count = len(pdf.pages)
                    for page_num, page in enumerate(pdf.pages):
                        text = page.extract_text()
                        if text:
                            resume_text += f"\n--- Page {page_num + 1} ---\n{text}"
            except Exception as e:
                logger.warning(f"PDF parsing warning: {e}")
                resume_text = "[PDF parsing issue, but file was read]"
                
        elif file_path.endswith('.docx'):
            try:
                doc = Document(file_path)
                page_count = len(doc.paragraphs)
                for para in doc.paragraphs:
                    if para.text.strip():
                        resume_text += f"{para.text}\n"
            except Exception as e:
                logger.warning(f"DOCX parsing warning: {e}")
                resume_text = "[DOCX parsing issue, but file was read]"
        else:
            raise ValueError("Unsupported file format. Use PDF or DOCX")
        
        if not resume_text.strip():
            raise ValueError("Could not extract text from resume")
        
        logger.info(f"✅ Resume parsed: {len(resume_text)} chars, {page_count} pages")
        
        return {
            "status": "success",
            "text": resume_text[:5000],  # Truncate for storage
            "full_text": resume_text,
            "pages": page_count,
            "task_id": self.request.id,
            "char_count": len(resume_text)
        }
        
    except Exception as exc:
        logger.error(f"❌ Resume parsing failed: {exc}")
        raise self.retry(exc=exc, countdown=60)


@shared_task(bind=True, max_retries=2)
def analyze_resume_async(self, resume_text: str, job_role: str):
    """
    Async task: Analyze resume using AI for skill matching
    Uses AI to extract skills, experience level, fit for role
    
    Args:
        resume_text: Extracted resume text
        job_role: Target job role
    
    Returns:
        dict: {"skills": [...], "experience_level": "...", "fit_score": 8.5}
    """
    try:
        from utils.ai_backends import get_ai_manager
        
        logger.info(f"Analyzing resume for role: {job_role}")
        
        ai = get_ai_manager()
        
        prompt = f"""Analyze this resume for a {job_role} position.
Resume:
{resume_text[:3000]}

Provide analysis in JSON format:
{{
  "skills": ["skill1", "skill2", ...],
  "experience_years": <number>,
  "experience_level": "entry|mid|senior",
  "fit_score": <1-10>,
  "strengths": "...",
  "gaps": "..."
}}

Return ONLY valid JSON."""
        
        analysis_json = ai.generate(prompt)
        
        # Parse JSON response
        import json
        try:
            analysis = json.loads(analysis_json)
        except:
            # Try to extract JSON from response
            import re
            match = re.search(r'\{.*\}', analysis_json, re.DOTALL)
            analysis = json.loads(match.group()) if match else {}
        
        logger.info(f"✅ Resume analyzed, fit score: {analysis.get('fit_score', 'N/A')}")
        
        return {
            "status": "success",
            "analysis": analysis,
            "task_id": self.request.id
        }
        
    except Exception as exc:
        logger.error(f"❌ Resume analysis failed: {exc}")
        raise self.retry(exc=exc, countdown=60)
