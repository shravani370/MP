from flask import Flask, redirect, request, render_template, session, url_for, flash, jsonify
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
import requests
import os
import tempfile
from dotenv import load_dotenv
from utils.ai_engine import generate_question, evaluate_answer
from utils.auth import hash_password, verify_password, require_csrf, set_secure_session, setup_secure_session, generate_csrf_token, validate_email, validate_password
from PyPDF2 import PdfReader
from docx import Document
from functools import wraps
import json
from sqlalchemy import select
from models.db import db, User, ScreeningResult, SavedJob, CoverLetter
from flask_mail import Mail

# Try to import Celery (optional - requires Python < 3.14)
try:
    from celery_app import app as celery_app
except (ModuleNotFoundError, ImportError):
    celery_app = None
    print("⚠️  WARNING: Celery not available (requires Python < 3.14). Async tasks disabled.")

# ================= ENV =================
load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "dev_key_change_in_production")

# Configure SQLAlchemy
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv(
    'DATABASE_URL',
    'postgresql://localhost/interview_proai'
)
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'pool_size': 10,
    'pool_recycle': 3600,
    'pool_pre_ping': True,
}
db.init_app(app)

# ═══════════════════════════════════════════════════════════════════════════
# REDIS CONFIGURATION (Sessions + Caching)
# ═══════════════════════════════════════════════════════════════════════════
REDIS_URL = os.getenv('REDIS_URL', 'redis://localhost:6379/0')

app.config.update(
    SESSION_TYPE='redis',
    SESSION_REDIS=REDIS_URL,
    SESSION_COOKIE_SECURE=os.getenv('SECURE_COOKIES', 'False').lower() in ('true', '1', 'yes'),
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE='Lax',
    PERMANENT_SESSION_LIFETIME=24 * 3600,  # 24 hours
)

# ═══════════════════════════════════════════════════════════════════════════
# EMAIL CONFIGURATION (Flask-Mail)
# ═══════════════════════════════════════════════════════════════════════════
app.config.update(
    MAIL_SERVER=os.getenv('MAIL_SERVER', 'localhost'),
    MAIL_PORT=int(os.getenv('MAIL_PORT', 25)),
    MAIL_USE_TLS=os.getenv('MAIL_USE_TLS', 'False').lower() in ('true', '1'),
    MAIL_USE_SSL=os.getenv('MAIL_USE_SSL', 'False').lower() in ('true', '1'),
    MAIL_USERNAME=os.getenv('MAIL_USERNAME', ''),
    MAIL_PASSWORD=os.getenv('MAIL_PASSWORD', ''),
    MAIL_DEFAULT_SENDER=os.getenv('MAIL_DEFAULT_SENDER', 'noreply@interview-proai.com')
)
mail = Mail(app)

# Initialize Celery with Flask app context (if available)
if celery_app:
    celery_app.conf.update(app.config)

    @celery_app.task(bind=True)
    def debug_task(self):
        print(f'Celery request: {self.request!r}')

# Setup secure session configuration
setup_secure_session(app)

# Setup rate limiting
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"],
    storage_uri="memory://"
)

# ── Enumerate filter for Jinja2 (used by screening templates) ──
app.jinja_env.filters['enumerate'] = enumerate

# ── Context processor for CSRF tokens in all templates ──
@app.context_processor
def inject_csrf_token():
    csrf_token = session.get('csrf_token', '')
    if not csrf_token:
        csrf_token = generate_csrf_token()
    return {"csrf_token": csrf_token}

# ═══════════════════════════════════════════════════════════════════════════
# DATABASE INITIALIZATION (use Alembic migrations in production)
# ═══════════════════════════════════════════════════════════════════════════
def init_db_with_app():
    """Initialize database tables (development only - use alembic in production)"""
    with app.app_context():
        try:
            db.create_all()
            print("✅ Database tables created/verified")
        except Exception as e:
            print(f"⚠️  Database initialization skipped: {str(e)}")

# Initialize DB on app start (ensures tables exist for development)
# Skip if database is not available
with app.app_context():
    try:
        db.create_all()
        print("✅ Database initialized successfully")
    except Exception as e:
        print(f"⚠️  Database not available, running in read-only mode: {str(e)[:100]}")


# ================= SCREENING BLUEPRINT =================
from screening.screening_routes import screening_bp
app.register_blueprint(screening_bp)

# ================= LOGIN REQUIRED =================
def login_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        if not session.get("user"):
            return redirect("/google-login")
        return f(*args, **kwargs)
    return wrapper

# ================= GOOGLE AUTH =================
CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID")
CLIENT_SECRET = os.getenv("GOOGLE_CLIENT_SECRET")
REDIRECT_URI = os.getenv("REDIRECT_URI", "http://127.0.0.1:5000/callback")

@app.route("/google-login")
def google_login():
    auth_url = (
        "https://accounts.google.com/o/oauth2/v2/auth?"
        f"client_id={CLIENT_ID}"
        f"&redirect_uri={REDIRECT_URI}"
        "&response_type=code"
        "&scope=openid email profile"
        "&access_type=offline"
        "&prompt=consent"
    )
    return redirect(auth_url)

@app.route("/callback")
def callback():
    code = request.args.get("code")
    if not code:
        return "Login failed ❌"
    try:
        token = requests.post(
            "https://oauth2.googleapis.com/token",
            data={
                "code": code,
                "client_id": CLIENT_ID,
                "client_secret": CLIENT_SECRET,
                "redirect_uri": REDIRECT_URI,
                "grant_type": "authorization_code"
            }
        ).json()

        access_token = token.get("access_token")
        user_info = requests.get(
            "https://www.googleapis.com/oauth2/v2/userinfo",
            headers={"Authorization": f"Bearer {access_token}"}
        ).json()

        name = user_info.get("name")
        email = user_info.get("email")

        # Check if user exists, create if not
        user = User.query.filter_by(email=email).first()
        if not user:
            user = User(
                name=name,
                email=email,
                password="google_auth",
                auth_type="google"
            )
            db.session.add(user)
            db.session.commit()

        set_secure_session(name, email)
        return redirect("/")

    except Exception as e:
        return f"Google Login Error: {e}"

# ================= HOME =================
@app.route("/")
def home():
    return render_template("index.html", user=session.get("user"))

@app.route("/dashboard")
def dashboard():
    if "user" not in session:
        return redirect(url_for("google_login"))
    
    email = session.get("email")
    user_name = session.get("user")
    
    # Fetch dashboard data
    try:
        # Get user from database
        user = User.query.filter_by(email=email).first()
        if not user:
            flash("❌ User not found", "error")
            return redirect(url_for("google_login"))
        
        # Get recent screening results (last 10)
        recent_results = db.session.query(ScreeningResult).filter_by(
            user_id=user.id
        ).order_by(ScreeningResult.created_at.desc()).limit(10).all()
        
        # Calculate statistics
        all_results = db.session.query(ScreeningResult).filter_by(user_id=user.id).all()
        total_attempts = len(all_results)
        passed_count = sum(1 for r in all_results if r.passed)
        
        if all_results:
            avg_score = sum((r.mcq_score and r.code_score and (r.mcq_score + r.code_score) / 2.0) or 0 
                           for r in all_results) / total_attempts if total_attempts > 0 else 0
            avg_score = round(avg_score, 2)
        else:
            avg_score = 0
        
        # Get saved jobs count
        saved_jobs_count = db.session.query(SavedJob).filter_by(user_id=user.id).count()
        
        # Format screening results for template
        screening_results = [
            {
                'role': r.role,
                'mcq_score': r.mcq_score,
                'code_score': r.code_score,
                'passed': r.passed,
                'created_at': r.created_at
            }
            for r in recent_results
        ]
        
        dashboard_data = {
            "user": user_name,
            "email": email,
            "total_attempts": total_attempts,
            "avg_score": avg_score,
            "passed_count": passed_count,
            "recent_screenings": screening_results,
            "saved_jobs": saved_jobs_count
        }
        
        return render_template("dashboard.html", **dashboard_data)
    except Exception as e:
        flash(f"Error loading dashboard: {e}", "error")
        return render_template("dashboard.html", user=user_name, error=str(e))

@app.route("/logout")
def logout():
    session.clear()
    return redirect("/")

# ================= LOGIN / SIGNUP / PROFILE =================
@app.route("/login", methods=["GET", "POST"])
@limiter.limit("5 per minute")
def login():
    if request.method == "POST":
        email = request.form.get("email", "").strip()
        password = request.form.get("password", "").strip()
        csrf_token = request.form.get("csrf_token")
        
        # Validate CSRF
        if not csrf_token or csrf_token != session.get("csrf_token"):
            flash("❌ Invalid request (CSRF validation failed)", "error")
            return render_template("login.html", csrf_token=generate_csrf_token())
        
        # Validate inputs
        if not email or not password:
            flash("❌ Email and password required", "error")
            return render_template("login.html", csrf_token=generate_csrf_token())
        
        try:
            # Fetch user from database
            user = User.query.filter_by(email=email, auth_type='email').first()
            
            if not user:
                flash("❌ Invalid email or password", "error")
                return render_template("login.html", csrf_token=generate_csrf_token())
            
            if not verify_password(password, user.password):
                flash("❌ Invalid email or password", "error")
                return render_template("login.html", csrf_token=generate_csrf_token())
            
            # Successful login
            set_secure_session(user.name, email)
            flash("✅ Logged in successfully!", "success")
            return redirect("/")
        
        except Exception as e:
            flash(f"❌ Login error: {e}", "error")
            return render_template("login.html", csrf_token=generate_csrf_token())
    
    # GET request
    return render_template("login.html", csrf_token=generate_csrf_token())


@app.route("/signup", methods=["GET", "POST"])
@limiter.limit("10 per hour")
def signup():
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        email = request.form.get("email", "").strip()
        password = request.form.get("password", "").strip()
        confirm_password = request.form.get("confirm_password", "").strip()
        csrf_token = request.form.get("csrf_token")
        
        # Validate CSRF
        if not csrf_token or csrf_token != session.get("csrf_token"):
            flash("❌ Invalid request (CSRF validation failed)", "error")
            return render_template("signup.html", csrf_token=generate_csrf_token())
        
        # Validate inputs
        if not name or not email or not password:
            flash("❌ Name, email, and password required", "error")
            return render_template("signup.html", csrf_token=generate_csrf_token())
        
        if not validate_email(email):
            flash("❌ Invalid email address", "error")
            return render_template("signup.html", csrf_token=generate_csrf_token())
        
        if password != confirm_password:
            flash("❌ Passwords do not match", "error")
            return render_template("signup.html", csrf_token=generate_csrf_token())
        
        # Validate password strength
        is_valid, msg = validate_password(password)
        if not is_valid:
            flash(f"❌ {msg}", "error")
            return render_template("signup.html", csrf_token=generate_csrf_token())
        
        try:
            # Check if email already exists
            existing_user = User.query.filter_by(email=email).first()
            if existing_user:
                flash("❌ Email already registered", "error")
                return render_template("signup.html", csrf_token=generate_csrf_token())
            
            # Create new user
            password_hash = hash_password(password)
            new_user = User(
                name=name,
                email=email,
                password=password_hash,
                auth_type='email'
            )
            db.session.add(new_user)
            db.session.commit()
            
            # Auto-login
            set_secure_session(name, email)
            flash("✅ Account created successfully!", "success")
            return redirect("/")
        
        except Exception as e:
            db.session.rollback()
            flash(f"❌ Signup error: {e}", "error")
            return render_template("signup.html", csrf_token=generate_csrf_token())
    
    # GET request
    return render_template("signup.html", csrf_token=generate_csrf_token())


@app.route("/profile", methods=["GET", "POST"])
def profile():
    if "user" not in session:
        return redirect(url_for("google_login"))
    
    email = session.get("email")
    user_name = session.get("user")
    
    if request.method == "POST":
        csrf_token = request.form.get("csrf_token")
        
        # Validate CSRF
        if not csrf_token or csrf_token != session.get("csrf_token"):
            flash("❌ Invalid request (CSRF validation failed)", "error")
            return render_template("profile.html", csrf_token=generate_csrf_token())
        
        # Handle profile update
        name = request.form.get("name", "").strip()
        
        if not name:
            flash("❌ Name is required", "error")
            return render_template("profile.html", csrf_token=generate_csrf_token())
        
        try:
            # Update user using ORM
            user = User.query.filter_by(email=email).first()
            if user:
                user.name = name
                db.session.commit()
                
                # Update session
                session["user"] = name
                session.modified = True
                
                flash("✅ Profile updated successfully!", "success")
                return redirect(url_for("profile"))
            else:
                flash("❌ User not found", "error")
                return render_template("profile.html", csrf_token=generate_csrf_token())
        
        except Exception as e:
            db.session.rollback()
            flash(f"❌ Error updating profile: {e}", "error")
            return render_template("profile.html", csrf_token=generate_csrf_token())
    
    # GET request - fetch user stats
    try:
        # Get user and their screening results using ORM
        user = User.query.filter_by(email=email).first()
        if not user:
            flash("❌ User not found", "error")
            return render_template("profile.html", user=user_name, csrf_token=generate_csrf_token())
        
        # Get all screening results for statistics
        all_results = db.session.query(ScreeningResult).filter_by(user_id=user.id).all()
        
        total_interviews = len(all_results)
        avg_score = 0
        if all_results:
            valid_scores = [r for r in all_results if r.mcq_score and r.code_score]
            if valid_scores:
                avg_score = round(sum((r.mcq_score + r.code_score) / 2.0 for r in valid_scores) / len(valid_scores), 2)
        
        # Get recent screenings (last 5)
        recent_interviews = db.session.query(ScreeningResult).filter_by(
            user_id=user.id
        ).order_by(ScreeningResult.created_at.desc()).limit(5).all()
        
        # Format results for template
        formatted_interviews = [
            {
                'role': r.role,
                'mcq_score': r.mcq_score,
                'code_score': r.code_score,
                'passed': r.passed,
                'created_at': r.created_at
            }
            for r in recent_interviews
        ]
        
        profile_data = {
            "user": user_name,
            "email": email,
            "total_interviews": total_interviews,
            "avg_score": avg_score,
            "recent_interviews": formatted_interviews,
            "csrf_token": generate_csrf_token()
        }
        
        return render_template("profile.html", **profile_data)
    
    except Exception as e:
        flash(f"Error loading profile: {e}", "error")
        return render_template("profile.html", user=user_name, csrf_token=generate_csrf_token())

# ================= INTERVIEW =================
@app.route("/start")
def start():
    if "user" not in session:
        return redirect(url_for("google_login"))

    raw_mode = request.args.get("mode", "")
    mode = raw_mode.strip().lower()
    if mode != "video":
        mode = "chat"

    topic = request.args.get("topic", "Machine Learning")
    role = request.args.get("role", topic)
    first_q = generate_question(role)

    # Preserve auth session, clear interview state
    current_user = session.get("user")
    current_email = session.get("email")
    session.clear()
    session["user"] = current_user
    session["email"] = current_email

    session["topic"] = role
    session["mode"] = mode
    session["question"] = first_q
    session["count"] = 0
    session["messages"] = [{"role": "ai", "text": first_q, "type": "question"}]
    session["asked_questions"] = [first_q]
    session["answers"] = []
    session["results"] = []

    if mode == "video":
        return render_template("video.html", question=first_q)
    return render_template("interview.html", messages=session["messages"])


@app.route("/interview", methods=["GET", "POST"])
def interview():
    """Handle interview page and form submission"""
    if request.method == "POST":
        return submit()
    if "user" not in session:
        return redirect(url_for("google_login"))
    messages = session.get("messages", [])
    if not messages:
        return redirect(url_for("start"))
    return render_template("interview.html", messages=messages)

@app.route("/submit", methods=["POST"])
def submit():
    answer = request.form.get("answer", "").strip()
    topic = session.get("topic")
    current_q = session.get("question")
    count = session.get("count", 0)
    mode = (session.get("mode") or "").strip().lower()
    if mode != "video":
        mode = "chat"

    messages = session.get("messages", [])
    asked = session.get("asked_questions", [])

    if not answer:
        if mode == "video":
            return render_template("video.html", question=current_q)
        return render_template("interview.html", messages=messages)

    # ── 1. Append candidate's answer ──
    messages.append({"role": "user", "text": answer})

    # ── 2. Initialise storage if missing ──
    if "answers" not in session:
        session["answers"] = []
    if "results" not in session:
        session["results"] = []

    session["answers"].append({"question": current_q, "answer": answer})

    # ── 3. Evaluate answer → get natural feedback ──
    result = evaluate_answer(current_q, answer)

    if isinstance(result, dict):
        session["results"].append(result)
        feedback_text = result.get("feedback", "").strip()
        score = result.get("score", 5)
    else:
        fallback_result = {
            "score": 5,
            "feedback": "Got it, thanks for that.",
            "strength": "Answer provided",
            "area_to_improve": "Add more specific examples"
        }
        session["results"].append(fallback_result)
        feedback_text = fallback_result["feedback"]
        score = 5

    # ── 4. Interviewer reacts to the answer (feedback bubble) ──
    if feedback_text:
        messages.append({
            "role": "ai",
            "text": feedback_text,
            "type": "feedback"
        })

    count += 1
    session["count"] = count

    # ── 5. If interview is done, go to results ──
    if count >= 5:
        session["messages"] = messages
        session.modified = True
        
        # Calculate average score
        total_score = sum(r.get("score", 0) for r in session["results"])
        avg_score = total_score / len(session["results"]) if session["results"] else 0
        avg_score = round(avg_score, 1)
        
        # Generate coaching tips
        coaching_prompt = f"""Based on these answers in a {topic} interview, provide 2-3 key improvement areas (brief bullet points):
Answers: {[a.get('answer', '') for a in session['answers']]}
Results: {session['results']}

Keep it concise and actionable."""
        
        try:
            from utils.ai_backends import get_ai_manager
            ai = get_ai_manager()
            coaching_tips = ai.generate(coaching_prompt)
        except:
            coaching_tips = "Keep practicing and review the feedback provided above."
        
        return render_template(
            "result.html",
            answers=session["answers"],
            results=session["results"],
            avg_score=avg_score,
            coaching_tips=coaching_tips,
            topic=topic,
            csrf_token=generate_csrf_token()
        )

    # ── 6. Generate the next question ──
    next_q = generate_question(
        topic,
        previous_answer=answer,
        history=messages,
        asked_questions=asked
    )

    # Deduplicate
    retry = 0
    while next_q in asked and retry < 3:
        next_q = generate_question(topic, previous_answer=answer, history=messages, asked_questions=asked)
        retry += 1

    if next_q in asked:
        next_q = f"What's been the most challenging aspect of working with {topic} for you?"

    # ── 7. Append the next question ──
    messages.append({
        "role": "ai",
        "text": next_q,
        "type": "question"
    })
    asked.append(next_q)

    session["messages"] = messages
    session["asked_questions"] = asked
    session["question"] = next_q
    session.modified = True

    if mode == "video":
        return render_template("video.html", question=next_q)
    return render_template("interview.html", messages=messages)

# ================= RESUME SUITE PAGE =================
@app.route("/resume_suite")
def resume_suite():
    if "user" not in session:
        return redirect(url_for("google_login"))
    # Redirect to template selection first
    return redirect(url_for("resume_templates"))

# ================= RESUME TEMPLATES =================
RESUME_TEMPLATES = {
    "classic": {
        "name": "Classic",
        "description": "Traditional, clean design. Perfect for formal industries.",
        "icon": "📄"
    },
    "modern": {
        "name": "Modern",
        "description": "Sleek and contemporary. Great for tech and creative roles.",
        "icon": "✨"
    },
    "professional": {
        "name": "Professional",
        "description": "Corporate style. Ideal for executive and business roles.",
        "icon": "💼"
    },
    "compact": {
        "name": "Compact",
        "description": "Space-efficient design. Fit more content on one page.",
        "icon": "📋"
    },
    "creative": {
        "name": "Creative",
        "description": "Colorful and bold. Perfect for designers and creatives.",
        "icon": "🎨"
    }
}

@app.route("/resume_templates")
def resume_templates():
    """Show available resume templates"""
    if "user" not in session:
        return redirect(url_for("google_login"))
    return render_template("resume_templates.html", templates=RESUME_TEMPLATES)

# ================= RESUME BUILDER =================
@app.route("/resume", methods=["GET", "POST"])
def resume():
    template = request.args.get("template", "classic")
    
    if template not in RESUME_TEMPLATES:
        template = "classic"
    
    if request.method == "POST":
        skills = [s.strip() for s in request.form.getlist("skills[]") if s.strip()]
        experience = [e.strip() for e in request.form.getlist("experience[]") if e.strip()]

        data = {
            "name": request.form.get("name", "").strip(),
            "role": request.form.get("role", "").strip(),
            "phone": request.form.get("phone", "").strip(),
            "email": request.form.get("email", "").strip(),
            "address": request.form.get("address", "").strip(),
            "about": request.form.get("about", "").strip(),
            "education": request.form.get("education", "").strip(),
            "experience": experience,
            "skills": skills,
            "template": template
        }

        if not data["name"] or not data["email"]:
            return render_template("resume.html", data=data, error="Name & Email required ❌", template=template, template_name=RESUME_TEMPLATES[template]["name"])

        return render_template("resume.html", data=data, success="Resume Generated ✅", template=template, template_name=RESUME_TEMPLATES[template]["name"])

    return render_template("resume.html", data={"skills": [], "experience": [], "template": template}, template=template, template_name=RESUME_TEMPLATES[template]["name"])

# ================= ANALYZER =================
@app.route("/analyze", methods=["POST"])
def analyze():
    resume_text = request.form.get("resume", "").strip()

    if not resume_text:
        return render_template("resume.html", result="❌ No resume content")

    try:
        from utils.ai_backends import get_ai_manager
        ai = get_ai_manager()
        prompt = f"""Analyze this resume and provide:
1. ATS Score (0-100)
2. Strengths (bullet points)
3. Weaknesses (bullet points)
4. Improvements (action items)

Resume:
{resume_text}"""
        result = ai.generate(prompt)
    except Exception as e:
        result = f"⚠️ AI error: {e}"

    return render_template("resume.html", result=result)

# ================= ANALYZE PAGE =================
@app.route("/analyze_page")
def analyze_page():
    if "user" not in session:
        return redirect(url_for("google_login"))
    return render_template("analyzer.html")

# ================= BUILDER PAGE =================
@app.route("/builder", methods=["GET", "POST"])
def builder():
    if "user" not in session:
        return redirect(url_for("google_login"))
    if request.method == "POST":
        return render_template("builder.html")
    return render_template("builder.html")

# ================= ATS =================
@app.route("/ats", methods=["GET", "POST"])
def ats():
    if request.method == "POST":
        file = request.files.get("resume_file")
        job_desc = request.form.get("job_desc", "").strip()

        if not file or file.filename == "":
            return render_template("ats.html", error="❌ No file uploaded")
        if not job_desc:
            return render_template("ats.html", error="❌ Job description required")
        if len(job_desc) < 50:
            return render_template("ats.html", error="❌ Job description too short (minimum 50 characters). Provide full details for accurate ATS analysis.")

        resume_text = ""
        filename = file.filename.lower()

        try:
            if filename.endswith(".pdf"):
                reader = PdfReader(file)
                for page in reader.pages:
                    resume_text += page.extract_text() or ""
            elif filename.endswith(".docx"):
                doc = Document(file)
                for para in doc.paragraphs:
                    resume_text += para.text + "\n"
            else:
                return render_template("ats.html", error="❌ Unsupported file type")
        except Exception as e:
            return render_template("ats.html", error=f"❌ File read error: {e}")

        if not resume_text.strip():
            return render_template("ats.html", error="❌ Could not extract text from file")

        try:
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".json", mode="w")
            json.dump({"resume_text": resume_text, "job_desc": job_desc}, tmp)
            tmp.close()
            session["ats_tmp"] = tmp.name
        except Exception as e:
            return render_template("ats.html", error=f"❌ Session error: {e}")

        return redirect(url_for("ats_result"))

    return render_template("ats.html")

# ================= ATS RESULT =================
@app.route("/ats_result")
def ats_result():
    tmp_path = session.get("ats_tmp")

    if not tmp_path or not os.path.exists(tmp_path):
        return redirect(url_for("ats"))

    try:
        with open(tmp_path, "r") as f:
            payload = json.load(f)
        resume_text = payload["resume_text"]
        job_desc = payload["job_desc"]
    except Exception:
        return redirect(url_for("ats"))
    finally:
        try:
            os.remove(tmp_path)
        except:
            pass
        session.pop("ats_tmp", None)

    try:
        from utils.ai_backends import get_ai_manager
        ai = get_ai_manager()
        
        prompt = f"""Return ONLY a JSON object. No explanation. No markdown. Just raw JSON.
You MUST fill ALL four fields.
- score: integer 0-100 based on resume matching job description
- matched_skills: list of skills in BOTH resume and job description
- missing_skills: list of skills in job description NOT in resume
- suggestions: list of exactly 3 specific actionable improvement tips

Format:
{{
  "score": <integer 0-100>,
  "matched_skills": ["skill1", "skill2"],
  "missing_skills": ["skill1", "skill2"],
  "suggestions": ["tip1", "tip2", "tip3"]
}}

JOB DESCRIPTION:
{job_desc}

RESUME:
{resume_text}

JSON:"""
        
        raw = ai.generate(prompt).strip()
        print("RAW AI RESPONSE:", raw)

        data = None
        
        # Attempt 1: Direct JSON parsing
        try:
            data = json.loads(raw)
            print("✅ Direct JSON parse successful")
        except Exception as e:
            print(f"❌ Direct parse failed: {e}")
        
        # Attempt 2: Strip markdown code blocks
        if not data:
            try:
                clean = raw.strip()
                if clean.startswith("```"):
                    clean = clean.split("```")[1]
                    if clean.startswith("json"):
                        clean = clean[4:]
                    clean = clean.strip()
                data = json.loads(clean)
                print("✅ Markdown strip parse successful")
            except Exception as e:
                print(f"❌ Markdown strip failed: {e}")
        
        # Attempt 3: Extract JSON from text
        if not data:
            try:
                start = raw.index("{")
                end = raw.rindex("}") + 1
                data = json.loads(raw[start:end])
                print("✅ Extract JSON parse successful")
            except Exception as e:
                print(f"❌ Extract JSON failed: {e}")
        
        # Attempt 4: Fix truncated JSON
        if not data:
            try:
                fixed = raw.strip()
                open_brackets = fixed.count("[") - fixed.count("]")
                open_braces = fixed.count("{") - fixed.count("}")
                for _ in range(open_brackets):
                    fixed += "]"
                for _ in range(open_braces):
                    fixed += "}"
                data = json.loads(fixed)
                print("✅ Fixed truncated JSON successful")
            except Exception as e:
                print(f"❌ Fix truncated failed: {e}")
        
        # Attempt 5: Alternative - if response looks like error, provide default
        if not data:
            if "[" in raw.lower() or "error" in raw.lower():
                print("⚠️  AI returned error or malformed response")
                data = {
                    "score": 25,
                    "matched_skills": [],
                    "missing_skills": [],
                    "suggestions": [
                        "AI analysis encountered an issue. Please try again.",
                        "Ensure your resume has clear skill listings.",
                        "Include industry keywords matching the job description."
                    ]
                }
            else:
                print("JSON PARSE FAILED. Raw was:", raw)
                data = {
                    "score": 0,
                    "matched_skills": [],
                    "missing_skills": [],
                    "suggestions": ["Could not parse AI response. Check terminal for raw output."]
                }

    except Exception as e:
        print("AI Error:", e)
        data = {
            "score": 0,
            "matched_skills": [],
            "missing_skills": [],
            "suggestions": [f"AI error: {e}"]
        }

    return render_template("ats_result.html", data=data)

# ================= JOB =================
@app.route("/job")
@login_required
def job():
    return render_template("job_role.html")

# ================= JOB SETUP =================
@app.route("/job_setup", methods=["GET", "POST"])
def job_setup():
    if "user" not in session:
        return redirect(url_for("google_login"))
    if request.method == "POST":
        session["job_role"] = request.form.get("role", "").strip()
        session["job_location"] = request.form.get("location", "").strip()
        job_types = request.form.getlist("job_type")
        session["job_types"] = job_types if job_types else ["Full-time", "Contract", "Part-time", "Internship"]
        return redirect(url_for("job_resume_page"))
    return render_template("job_role.html")

# ================= JOB RESUME UPLOAD PAGE =================
@app.route("/job_resume_page")
@app.route("/job_resume")
def job_resume_page():
    if "user" not in session:
        return redirect(url_for("google_login"))
    if not session.get("job_role"):
        return redirect(url_for("job"))
    return render_template("job_resume.html")

# ================= COUNTRY MAP =================
COUNTRY_MAP = {
    "india": "in", "us": "us", "usa": "us", "united states": "us",
    "uk": "gb", "united kingdom": "gb", "australia": "au", "canada": "ca",
    "germany": "de", "france": "fr", "netherlands": "nl", "singapore": "sg",
    "new zealand": "nz", "south africa": "za", "brazil": "br", "russia": "ru",
    "poland": "pl", "austria": "at", "belgium": "be", "switzerland": "ch",
    "italy": "it", "mexico": "mx", "spain": "es", "argentina": "ar",
}


def detect_country_code(location: str) -> tuple[str, str]:
    loc_lower = location.lower()
    country_code = "in"
    for country_name, code in COUNTRY_MAP.items():
        if country_name in loc_lower:
            country_code = code
            break
    city_location = location
    for country_name in COUNTRY_MAP:
        city_location = city_location.lower().replace(country_name, "").strip(" ,")
    return country_code, city_location


def extract_skills_from_resume(file) -> str:
    resume_text = ""
    if not file or not file.filename:
        return ""
    filename = file.filename.lower()
    try:
        if filename.endswith(".pdf"):
            reader = PdfReader(file)
            for page in reader.pages:
                resume_text += page.extract_text() or ""
        elif filename.endswith(".docx"):
            doc = Document(file)
            for para in doc.paragraphs:
                resume_text += para.text + "\n"
        else:
            resume_text = file.read().decode("utf-8", errors="ignore")
    except Exception as e:
        print("Resume read error:", e)
        return ""
    if not resume_text.strip():
        return ""
    try:
        from utils.ai_backends import get_ai_manager
        ai = get_ai_manager()
        prompt = f"""Extract a comma-separated list of technical skills from this resume.
Return ONLY the skills, nothing else (no markdown, no explanation):

{resume_text}"""
        return ai.generate(prompt).strip()
    except Exception as e:
        print("AI skill extraction error:", e)
        return ""


def search_adzuna_jobs(country_code: str, query: str, location: str = "", page: int = 1, job_types: list = None) -> list:
    APP_ID = os.getenv("APP_ID")
    APP_KEY = os.getenv("APP_KEY")

    search_query = query
    if job_types and len(job_types) == 1 and "Internship" in job_types:
        search_query = f"{query} internship"
    elif job_types and "Internship" in job_types and len(job_types) > 1:
        search_query = f"{query} internship"

    params = {
        "app_id": APP_ID,
        "app_key": APP_KEY,
        "what": search_query,
        "results_per_page": 10,
        "sort_by": "relevance",
        "content-type": "application/json",
    }
    if location:
        params["where"] = location

    url = f"https://api.adzuna.com/v1/api/jobs/{country_code}/search/{page}"
    print(f"Adzuna → {url} | params: {params}")

    response = requests.get(url, params=params, timeout=15)
    response.raise_for_status()
    results = response.json().get("results", [])
    
    # Map ADZUNA response fields to template expectations
    for job in results:
        # Rename redirect_url to url for template compatibility
        if "redirect_url" in job and "url" not in job:
            job["url"] = job["redirect_url"]
        
        # Extract company display_name
        if isinstance(job.get("company"), dict) and "display_name" in job["company"]:
            job["company"] = job["company"]["display_name"]
        
        # Extract location display_name
        if isinstance(job.get("location"), dict) and "display_name" in job["location"]:
            job["location"] = job["location"]["display_name"]
        
        # Extract salary if available
        if "salary_max" in job and job["salary_max"]:
            job["salary"] = job["salary_max"]
        elif "salary_min" in job and job["salary_min"]:
            job["salary"] = job["salary_min"]
    
    return results


# ================= SAVED JOBS =================
@app.route("/saved_jobs")
@login_required
def saved_jobs():
    email = session.get("email")
    if not email:
        return redirect(url_for("google_login"))
    
    try:
        # Get user and their saved jobs using ORM
        user = User.query.filter_by(email=email).first()
        if not user:
            jobs = []
        else:
            saved_jobs_query = db.session.query(SavedJob).filter_by(
                user_id=user.id
            ).order_by(SavedJob.saved_at.desc()).all()
            
            # Convert to dictionaries for template compatibility
            jobs = [
                {
                    'id': job.id,
                    'job_id': str(job.job_id),
                    'title': job.title,
                    'company': job.company,
                    'location': job.location,
                    'url': job.url,
                    'saved_at': job.saved_at
                }
                for job in saved_jobs_query
            ]
    except Exception as e:
        print(f"Database error in saved_jobs: {e}")
        jobs = []
    
    return render_template("saved_jobs.html", jobs=jobs)

@app.route("/api/save_job", methods=["POST"])
def api_save_job():
    """API endpoint to save or remove a job"""
    if "email" not in session:
        return {"status": "error", "message": "Not logged in"}, 401
    
    try:
        data = request.get_json()
        job_id = data.get("job_id")
        action = data.get("action")  # 'save' or 'unsave'
        email = session.get("email")
        
        if not job_id:
            return {"status": "error", "message": "Missing job_id"}, 400
        
        try:
            # Get user
            user = User.query.filter_by(email=email).first()
            if not user:
                return {"status": "error", "message": "User not found"}, 404
            
            if action == "save":
                # Check if job already saved
                existing = SavedJob.query.filter_by(
                    user_id=user.id,
                    job_id=job_id
                ).first()
                
                if not existing:
                    # Create new saved job
                    saved_job = SavedJob(
                        user_id=user.id,
                        email=email,
                        job_id=job_id,
                        title=data.get("title", ""),
                        company=data.get("company", ""),
                        location=data.get("location", ""),
                        url=data.get("url", "")
                    )
                    db.session.add(saved_job)
                    db.session.commit()
                    return {"status": "ok", "message": "Job saved successfully"}
                else:
                    return {"status": "ok", "message": "Job already saved"}
                    
            elif action == "unsave":
                # Delete saved job
                saved_job = SavedJob.query.filter_by(
                    user_id=user.id,
                    job_id=job_id
                ).first()
                
                if saved_job:
                    db.session.delete(saved_job)
                    db.session.commit()
                    return {"status": "ok", "message": "Job unsaved successfully"}
                else:
                    return {"status": "ok", "message": "Job was not saved"}
            else:
                return {"status": "error", "message": "Invalid action"}, 400
                
        except Exception as db_error:
            db.session.rollback()
            print(f"Database error in api_save_job: {db_error}")
            return {"status": "error", "message": str(db_error)}, 500
    except Exception as e:
        print(f"Error in api_save_job: {e}")
        return {"status": "error", "message": str(e)}, 500

# ================= SKILLS GAP ANALYSIS =================
@app.route("/skills", methods=["GET", "POST"])
@login_required
def skills():
    """Skill gap analysis page"""
    if request.method == "POST":
        role = request.form.get("role", "").strip()
        user_skills = request.form.get("user_skills", "").strip()
        
        if not role or not user_skills:
            return render_template("skills.html", roles=["Python Developer", "Data Scientist", "DevOps Engineer", "Full Stack Developer", "Frontend Developer", "Backend Developer", "Mobile Developer", "QA Engineer", "Cloud Architect", "Machine Learning Engineer"], gap_result=None)
        
        # Skills required for different roles (predefined mapping)
        role_skills_map = {
            "Python Developer": ["Python", "Django", "FastAPI", "SQL", "Git", "Docker", "REST APIs", "Unit Testing"],
            "Data Scientist": ["Python", "Pandas", "NumPy", "Scikit-learn", "TensorFlow", "SQL", "Statistics", "Matplotlib"],
            "DevOps Engineer": ["Docker", "Kubernetes", "Jenkins", "Git", "Linux", "AWS", "Terraform", "CI/CD"],
            "Full Stack Developer": ["JavaScript", "React", "Node.js", "MongoDB", "SQL", "HTML", "CSS", "Git"],
            "Frontend Developer": ["JavaScript", "React", "Vue.js", "CSS", "HTML", "Webpack", "Git", "UI/UX"],
            "Backend Developer": ["Node.js", "Python", "Java", "SQL", "APIs", "Docker", "Git", "Database Design"],
            "Mobile Developer": ["React Native", "Flutter", "Swift", "Kotlin", "Firebase", "Git", "REST APIs", "Mobile UI"],
            "QA Engineer": ["Selenium", "Jest", "Cypress", "Manual Testing", "SQL", "Git", "JIRA", "Test Planning"],
            "Cloud Architect": ["AWS", "Azure", "GCP", "Docker", "Kubernetes", "Terraform", "Networking", "Security"],
            "Machine Learning Engineer": ["Python", "TensorFlow", "PyTorch", "Scikit-learn", "SQL", "Statistics", "Git", "Data Analysis"],
        }
        
        required = role_skills_map.get(role, [])
        user_skills_list = [s.strip() for s in user_skills.split(",")]
        matched = [s for s in required if any(req.lower() in s.lower() or s.lower() in req.lower() for req in user_skills_list)]
        missing = [s for s in required if s not in matched]
        match_pct = int(len(matched) / len(required) * 100) if required else 0
        
        # Generate learning resources for missing skills
        resources = {}
        for skill in missing:
            skill_encoded = skill.replace(" ", "+")
            resources[skill] = {
                "coursera": f"https://www.coursera.org/search?query={skill_encoded}",
                "youtube": f"https://www.youtube.com/results?search_query={skill_encoded}+tutorial"
            }
        
        gap_result = {
            "role": role,
            "required": required,
            "matched": matched,
            "missing": missing,
            "match_pct": match_pct,
            "resources": resources
        }
        return render_template("skills.html", roles=["Python Developer", "Data Scientist", "DevOps Engineer", "Full Stack Developer", "Frontend Developer", "Backend Developer", "Mobile Developer", "QA Engineer", "Cloud Architect", "Machine Learning Engineer"], gap_result=gap_result, prefill_skills=user_skills)
    
    return render_template("skills.html", roles=["Python Developer", "Data Scientist", "DevOps Engineer", "Full Stack Developer", "Frontend Developer", "Backend Developer", "Mobile Developer", "QA Engineer", "Cloud Architect", "Machine Learning Engineer"], gap_result=None, prefill_skills="")

# ================= JOB RESULT =================
@app.route("/job_result", methods=["GET", "POST"])
def job_result():
    role = session.get("job_role", "")
    location = session.get("job_location", "")
    job_types = session.get("job_types", ["Full-time", "Contract", "Part-time", "Internship"])

    skills = ""
    if request.method == "POST":
        skills = extract_skills_from_resume(request.files.get("resume"))

    country_code, city_location = detect_country_code(location)
    search_query = role.strip() if role.strip() else "software engineer"

    jobs = []
    error_msg = None

    try:
        jobs = search_adzuna_jobs(country_code, search_query, city_location, job_types=job_types)

        if not jobs and city_location:
            print("No results with location — retrying without location filter")
            jobs = search_adzuna_jobs(country_code, search_query, job_types=job_types)
            if jobs:
                error_msg = f"No jobs found in '{location}' — showing remote / nationwide results instead."

        if not jobs:
            broad_query = role.split()[0] if role.strip() else "developer"
            print(f"Still no results — retrying with broad query: '{broad_query}'")
            jobs = search_adzuna_jobs(country_code, broad_query, job_types=job_types)
            if jobs:
                error_msg = f"No exact matches for '{role}' — showing broader results for '{broad_query}'."

        if not jobs and country_code != "gb":
            print("Falling back to GB endpoint")
            jobs = search_adzuna_jobs("gb", search_query, job_types=job_types)
            if jobs:
                error_msg = "No results found in your region — showing international listings instead."

        if not jobs:
            error_msg = "No jobs found. Try a broader role title or a different location."

    except requests.exceptions.HTTPError as e:
        print(f"Adzuna HTTP error: {e}")
        error_msg = f"Job search API returned an error: {e}"
    except requests.exceptions.ConnectionError:
        error_msg = "Could not connect to job search service. Check your internet connection."
    except requests.exceptions.Timeout:
        error_msg = "Job search request timed out. Please try again."
    except Exception as e:
        print("Job API unexpected error:", e)
        error_msg = f"Unexpected error during job search: {e}"

    session.pop("job_role", None)
    session.pop("job_location", None)
    session.pop("job_types", None)

    # Add saved status to each job
    email = session.get("email")
    if email:
        try:
            user = User.query.filter_by(email=email).first()
            if user:
                saved_jobs_query = db.session.query(SavedJob.job_id).filter_by(user_id=user.id).all()
                saved_job_ids = {str(row[0]) for row in saved_jobs_query}
                for job in jobs:
                    job["saved"] = str(job.get("id", "")) in saved_job_ids
            else:
                for job in jobs:
                    job["saved"] = False
        except Exception as db_error:
            print(f"Database error in job_result: {db_error}")
            # Mark all jobs as not saved if query fails
            for job in jobs:
                job["saved"] = False
    else:
        for job in jobs:
            job["saved"] = False
    
    # Create job saves dict for template
    job_saves = {job.get("id", ""): job.get("saved", False) for job in jobs}
    
    return render_template(
        "job_result.html",
        jobs=jobs,
        jobs_saves=job_saves,
        skills=skills,
        role=role,
        location=location,
        job_types=job_types,
        error_msg=error_msg,
    )

# ================= COVER LETTER =================
@app.route("/cover_letter", methods=["GET", "POST"])
@login_required
def cover_letter():
    email = session.get("email")
    if not email:
        return redirect(url_for("google_login"))
    
    letter = None
    saved = []
    
    # Get user (needed for all operations)
    user = User.query.filter_by(email=email).first()
    if not user:
        flash("❌ User not found", "error")
        return redirect(url_for("google_login"))
    
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        role = request.form.get("role", "").strip()
        company = request.form.get("company", "").strip()
        job_desc = request.form.get("job_desc", "").strip()
        resume_text = request.form.get("resume_text", "").strip()
        
        if not name or not role or not company:
            flash("❌ Name, role, and company are required", "error")
            return render_template("cover_letter.html", letter=None, saved=[])
        
        if not job_desc:
            flash("❌ Job description is required", "error")
            return render_template("cover_letter.html", letter=None, saved=[])
        
        try:
            # Generate cover letter using AI
            from utils.ai_backends import get_ai_manager
            ai = get_ai_manager()
            
            prompt = f"""Write a professional cover letter for:
- Name: {name}
- Role: {role}
- Company: {company}
- Key job requirements: {job_desc}
{f'- My experience: {resume_text}' if resume_text else ''}

Write a compelling, personalized 3-4 paragraph cover letter. No salutation/date needed, just the body."""
            
            letter = ai.generate(prompt).strip()
            
            if not letter:
                flash("❌ AI failed to generate letter. Try again.", "error")
                return render_template("cover_letter.html", letter=None, saved=[])
            
            # Save to database using ORM
            try:
                cover_letter_obj = CoverLetter(
                    user_id=user.id,
                    email=email,
                    name=name,
                    role=role,
                    company=company,
                    job_desc=job_desc,
                    resume_text=resume_text,
                    letter=letter
                )
                db.session.add(cover_letter_obj)
                db.session.commit()
                flash("✅ Cover letter saved!", "success")
            except Exception as save_error:
                db.session.rollback()
                print(f"Database error in cover_letter: {save_error}")
                # Letter is still generated even if save fails
                flash("⚠️ Letter generated but not saved", "warning")
        
        except requests.exceptions.Timeout:
            flash("❌ AI request timed out. Check if Ollama is running.", "error")
        except requests.exceptions.ConnectionError:
            flash("❌ Cannot connect to AI service. Check Ollama is running.", "error")
        except Exception as e:
            print(f"Error in cover_letter generation: {e}")
            flash(f"❌ Error: {str(e)}", "error")
    
    # Load saved letters
    try:
        saved_letters = db.session.query(CoverLetter).filter_by(
            user_id=user.id
        ).order_by(CoverLetter.created_at.desc()).limit(10).all()
        
        saved = [
            {
                'id': cl.id,
                'name': cl.name,
                'role': cl.role,
                'company': cl.company,
                'created_at': cl.created_at
            }
            for cl in saved_letters
        ]
    except Exception as e:
        print(f"Database error loading saved letters: {e}")
        saved = []
    
    return render_template("cover_letter.html", letter=letter, saved=saved)

# ================= RUN =================
@app.route("/health")
def health_check():
    """Health check endpoint for monitoring and Docker"""
    try:
        # Check database connection
        db.session.execute(select(1))
        db_status = "🟢 OK"
    except Exception as e:
        db_status = f"🔴 ERROR: {str(e)}"
    
    return jsonify({
        "status": "healthy" if "OK" in db_status else "unhealthy",
        "database": db_status,
        "version": "1.0.0"
    }), 200 if "OK" in db_status else 503


# ═══════════════════════════════════════════════════════════════════════════
# ASYNC TASK API ENDPOINTS (For frontend polling)
# ═══════════════════════════════════════════════════════════════════════════

@app.route("/api/task/<task_id>")
def get_task_status(task_id):
    """
    Get status of an async task
    
    Response:
    {
      "task_id": "abc123",
      "state": "PENDING|PROGRESS|SUCCESS|FAILURE",
      "result": {...},
      "error": "optional error message"
    }
    """
    if not celery_app:
        return jsonify({"error": "Async tasks not available (Celery not installed)"}), 503
    
    try:
        result = celery_app.AsyncResult(task_id)
        
        response = {
            "task_id": task_id,
            "state": result.state,
        }
        
        if result.state == 'SUCCESS':
            response["result"] = result.result
        elif result.state == 'FAILURE':
            response["error"] = str(result.info)
        elif result.state == 'PROGRESS':
            response["progress"] = result.info.get('current', 0)
            response["total"] = result.info.get('total', 100)
        
        return jsonify(response)
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/generate-question", methods=["POST"])
@login_required
def api_generate_question():
    """
    Submit async task to generate interview question
    Returns task ID for polling
    """
    try:
        from tasks.ai_tasks import generate_question_async
        
        data = request.get_json()
        topic = data.get("topic", "General")
        previous_answer = data.get("previous_answer", "")
        history = data.get("history", [])
        asked_questions = data.get("asked_questions", [])
        
        # Submit task
        task = generate_question_async.delay(
            topic=topic,
            previous_answer=previous_answer,
            history=history,
            asked_questions=asked_questions
        )
        
        return jsonify({
            "task_id": task.id,
            "status": "queued"
        }), 202
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/evaluate-answer", methods=["POST"])
@login_required
def api_evaluate_answer():
    """
    Submit async task to evaluate answer
    Returns task ID for polling
    """
    try:
        from tasks.ai_tasks import evaluate_answer_async
        
        data = request.get_json()
        question = data.get("question", "")
        answer = data.get("answer", "")
        
        if not question or not answer:
            return jsonify({"error": "Missing question or answer"}), 400
        
        # Submit task
        task = evaluate_answer_async.delay(
            question=question,
            answer=answer
        )
        
        return jsonify({
            "task_id": task.id,
            "status": "queued"
        }), 202
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/parse-resume", methods=["POST"])
@login_required
def api_parse_resume():
    """
    Submit async task to parse resume
    Returns task ID for polling
    """
    try:
        from tasks.resume_tasks import parse_resume_async
        
        email = session.get("email")
        file = request.files.get("resume")
        
        if not file:
            return jsonify({"error": "No file provided"}), 400
        
        # Save file temporarily
        import tempfile
        tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf' if file.filename.endswith('.pdf') else '.docx')
        file.save(tmp_file.name)
        
        # Submit task
        task = parse_resume_async.delay(
            file_path=tmp_file.name,
            user_email=email
        )
        
        return jsonify({
            "task_id": task.id,
            "status": "queued"
        }), 202
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/send-email", methods=["POST"])
@login_required
def api_send_email():
    """
    Submit async task to send email
    """
    try:
        from tasks.email_tasks import send_screening_results_email
        
        email = session.get("email")
        data = request.get_json()
        
        role = data.get("role", "")
        mcq_score = data.get("mcq_score", 0)
        code_score = data.get("code_score", 0)
        passed = data.get("passed", False)
        
        # Submit task
        task = send_screening_results_email.delay(
            user_email=email,
            role=role,
            mcq_score=mcq_score,
            code_score=code_score,
            passed=passed
        )
        
        return jsonify({
            "task_id": task.id,
            "status": "queued",
            "message": "Email will be sent shortly"
        }), 202
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ═══════════════════════════════════════════════════════════════════════════
# CELERY WORKER MONITORING
# ═══════════════════════════════════════════════════════════════════════════

@app.route("/admin/workers")
def admin_workers():
    """
    Show Celery worker status (admin only for now)
    In production, add proper auth
    """
    if not celery_app:
        return jsonify({"error": "Celery not available (requires Python < 3.14)", "active_tasks": {}}), 503
    
    try:
        # Get active tasks
        active_tasks = celery_app.control.inspect().active()
        scheduled_tasks = celery_app.control.inspect().scheduled()
        registered_tasks = celery_app.control.inspect().registered()
        
        return jsonify({
            "active_tasks": active_tasks or {},
            "scheduled_tasks": scheduled_tasks or {},
            "registered_tasks": registered_tasks or {}
        })
        
    except Exception as e:
        return jsonify({"error": str(e), "message": "Celery workers not responding"}), 500


if __name__ == "__main__":
    app.run(debug=True)